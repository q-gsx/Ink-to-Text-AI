import os
import streamlit as st
import requests
import base64
from PIL import Image, ImageOps, ImageFilter, ImageEnhance
import io
import time
import re
import json
import zipfile
import hashlib
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
from arabic_reshaper import reshape
from bidi.algorithm import get_display

# ============================================================
# 1. PAGE CONFIG — MUST BE FIRST
# ============================================================
st.set_page_config(
    page_title="Ink to text AI — Handwriting & Print OCR",
    page_icon="https://i.ibb.co/V0TcJzmL/LOGO2.jpg",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================
# 2. CSS + FONT AWESOME INJECTION
# ============================================================
from css_block import LIGHT_CSS, DARK_CSS

def _inject_css():
    is_dark = st.session_state.get("dark_mode", False)
    payload  = LIGHT_CSS
    if is_dark:
        payload += DARK_CSS
    if hasattr(st, "html"):
        st.html(payload)
    else:
        st.markdown(payload, unsafe_allow_html=True)

_inject_css()

# 3. CONSTANTS & CONFIG
# ============================================================
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
except Exception:
    API_KEY = "PLEASE_SET_YOUR_API_KEY_IN_SECRETS"

GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
MAX_IMAGE_SIZE_MB = 10
SUPPORTED_FORMATS = ['png', 'jpg', 'jpeg', 'webp', 'bmp', 'tiff', 'tif']

# ============================================================
# 4. SESSION STATE INIT
# ============================================================
defaults_state = {
    "extracted_result":    None,
    "processing_history":  [],
    "total_processed":     0,
    "total_chars":         0,
    "current_image_meta":  {},
    "enhancement_applied": False,
    "last_process_time":   0.0,
    "batch_results":       [],
    "notes":               "",
    "show_raw":            False,
    "dark_mode":           False,
}
for k, v in defaults_state.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ============================================================
# 5. UTILITY FUNCTIONS
# ============================================================
def file_size_mb(file_obj) -> float:
    file_obj.seek(0, 2)
    size = file_obj.tell() / (1024 * 1024)
    file_obj.seek(0)
    return round(size, 2)

def image_to_b64(pil_img: Image.Image, quality: int = 85) -> str:
    buf = io.BytesIO()
    pil_img.convert("RGB").save(buf, format="JPEG", quality=quality)
    return base64.b64encode(buf.getvalue()).decode("utf-8")

def clean_html(text: str) -> str:
    text = text.replace('```html', '').replace('```', '')
    text = re.sub(r'>\s+<', '><', text)
    return text.strip()

def strip_tags(text: str) -> str:
    return re.sub(r'<[^>]+>', '', text).strip()

def count_words(text: str) -> int:
    return len(strip_tags(text).split())

def count_chars(text: str) -> int:
    return len(strip_tags(text))

def detect_language(text: str) -> str:
    arabic_chars  = len(re.findall(r'[\u0600-\u06FF]', text))
    english_chars = len(re.findall(r'[a-zA-Z]', text))
    if arabic_chars > english_chars:
        return "عربي"
    elif english_chars > arabic_chars:
        return "English"
    return "Mixed"

def detect_content_type(text: str) -> str:
    if "<td>" in text.lower():
        return "جدول + نص"
    digits = len(re.findall(r'\d', text))
    if digits > 30:
        return "بيانات رقمية"
    return "نص عادي"

def file_hash(file_obj) -> str:
    file_obj.seek(0)
    h_full = hashlib.md5(file_obj.read()).hexdigest()
    h = h_full[0:8]
    file_obj.seek(0)
    return h

def enhance_image(pil_img: Image.Image,
                  contrast: float = 1.3,
                  sharpness: float = 1.4,
                  brightness: float = 1.05,
                  denoise: bool = False) -> Image.Image:
    img = pil_img.copy()
    if brightness != 1.0:
        img = ImageEnhance.Brightness(img).enhance(brightness)
    if contrast != 1.0:
        img = ImageEnhance.Contrast(img).enhance(contrast)
    if sharpness != 1.0:
        img = ImageEnhance.Sharpness(img).enhance(sharpness)
    if denoise:
        img = img.filter(ImageFilter.MedianFilter(size=3))
    return img

def resize_for_api(pil_img: Image.Image, max_px: int = 2048) -> Image.Image:
    w, h = pil_img.size
    if max(w, h) > max_px:
        scale = max_px / max(w, h)
        pil_img = pil_img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    return pil_img

def shape_arabic(text: str) -> str:
    """Reshape and apply BiDi to a string for correct Arabic PDF rendering."""
    try:
        return get_display(reshape(text))
    except Exception:
        return text

# ============================================================
# 6. EXPORT FUNCTIONS
# ============================================================
@st.cache_data(show_spinner=False)
def create_txt(content: str) -> bytes:
    clean = strip_tags(clean_html(content))
    return clean.encode("utf-8")

@st.cache_data(show_spinner=False)
def create_html_export(content: str, title: str = "DocuVision Export") -> bytes:
    html = f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<link rel="stylesheet"
      href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.0/css/all.min.css"
      crossorigin="anonymous" />
<style>
  body {{ font-family:'Segoe UI',Tahoma,'Noto Sans Arabic',Arial,sans-serif;
         background:#f8fafc; color:#1e293b; padding:3rem;
         max-width:860px; margin:auto; direction:rtl; }}
  h1   {{ color:#4f46e5; font-size:1.4rem; border-bottom:2px solid #e2e8f0;
         padding-bottom:0.5rem; margin-bottom:1.5rem; }}
  table{{ width:100%; border-collapse:collapse; margin:1rem 0; }}
  th   {{ background:#4f46e5; color:#fff; padding:10px; text-align:right; }}
  td   {{ border:1px solid #e2e8f0; padding:8px 12px; }}
  tr:nth-child(even) td {{ background:#f1f5f9; }}
  .meta{{ color:#64748b; font-size:0.82rem; margin-bottom:1.5rem; }}
  i    {{ margin-left:0.4rem; }}
</style>
</head>
<body>
<h1><i class="fa-solid fa-file-lines"></i> {title}</h1>
<p class="meta">
  <i class="fa-solid fa-robot"></i> Generated by Ink to text AI &mdash;
  {datetime.now().strftime('%Y-%m-%d %H:%M')}
</p>
<div>{clean_html(content)}</div>
</body>
</html>"""
    return html.encode("utf-8")

@st.cache_data(show_spinner=False)
def create_word_doc(content: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.right_to_left = True
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    title_para = doc.add_heading("", level=0)
    title_run  = title_para.add_run("Ink to text AI — Extracted Text")
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(79, 70, 229)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

    cleaned = clean_html(content)

    if "</table>" in cleaned.lower():
        parts = re.split(r'(<table.*?</table>)', cleaned, flags=re.DOTALL | re.IGNORECASE)
        for part in parts:
            if re.match(r'<table', part, re.IGNORECASE):
                rows_html = re.findall(r'<tr>(.*?)</tr>', part, re.DOTALL | re.IGNORECASE)
                if not rows_html:
                    continue
                max_cols = max(
                    len(re.findall(r'<(?:td|th)[^>]*>(.*?)</(?:td|th)>', r, re.DOTALL | re.IGNORECASE))
                    for r in rows_html
                )
                if max_cols == 0:
                    continue
                tbl = doc.add_table(rows=len(rows_html), cols=max_cols)
                tbl.style = 'Table Grid'
                for ri, row_html in enumerate(rows_html):
                    cells_html = re.findall(
                        r'<(?:td|th)[^>]*>(.*?)</(?:td|th)>', row_html, re.DOTALL | re.IGNORECASE
                    )
                    for ci, cell_txt in enumerate(cells_html):
                        if ci < max_cols:
                            cell = tbl.cell(ri, ci)
                            cell.text = strip_tags(cell_txt).strip()
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(79, 70, 229)
                doc.add_paragraph()
            else:
                text = strip_tags(part).strip()
                if text:
                    p   = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(text)
                    run.font.size = Pt(12)
    else:
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(strip_tags(cleaned))
        run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

class ArabicPDF(FPDF):
    """FPDF subclass with Arabic helper utilities."""

    FONT_REGULAR = "ArabicFont"
    INDIGO = (79, 70, 229)
    SLATE  = (100, 116, 139)
    DARK   = (30,  41,  59)
    LIGHT_BG = (248, 250, 252)
    WHITE    = (255, 255, 255)
    BORDER   = (226, 232, 240)

    def __init__(self, font_path: str | None = None):
        super().__init__('P', 'mm', 'A4')
        self.set_auto_page_break(auto=True, margin=20)
        self.set_margins(left=20, top=25, right=20)
        self._font_path = font_path
        self._font_loaded = False
        self._add_arabic_font()

    def _add_arabic_font(self):
        fp = self._font_path
        if fp is not None and os.path.exists(fp):
            try:
                self.add_font(self.FONT_REGULAR, '', fp, uni=True)
                self._font_loaded = True
            except Exception:
                self._font_loaded = False

    def set_arabic_font(self, size: int = 12):
        if self._font_loaded:
            self.set_font(self.FONT_REGULAR, '', size)
        else:
            self.set_font('Helvetica', '', size)

    def header(self):
        # Coloured header bar
        self.set_fill_color(*self.INDIGO)
        self.rect(0, 0, 210, 12, 'F')
        self.set_y(2)
        self.set_text_color(*self.WHITE)
        self.set_arabic_font(8)
        stamp = f"Ink to text AI  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        self.cell(0, 8, shape_arabic(stamp), align='C')
        self.set_text_color(*self.DARK)
        self.set_y(16)

    def footer(self):
        self.set_y(-12)
        self.set_draw_color(*self.BORDER)
        self.line(20, self.get_y(), 190, self.get_y())
        self.set_y(-10)
        self.set_text_color(*self.SLATE)
        self.set_arabic_font(7)
        page_text = shape_arabic(f"صفحة {self.page_no()}")
        self.cell(0, 5, page_text, align='C')

    def section_title(self, title: str):
        """Draw a shaded section heading."""
        self.set_fill_color(238, 242, 255)     # indigo-50
        self.set_draw_color(*self.INDIGO)
        self.set_text_color(*self.INDIGO)
        self.set_arabic_font(13)
        self.set_line_width(0.4)
        self.cell(0, 9, shape_arabic(title), border='B', ln=True, fill=True, align='R')
        self.set_line_width(0.2)
        self.ln(2)
        self.set_text_color(*self.DARK)

    def body_line(self, text: str, line_height: float = 7):
        """Write a single shaped Arabic line."""
        self.set_arabic_font(11)
        self.set_text_color(*self.DARK)
        shaped = shape_arabic(text)
        self.multi_cell(0, line_height, shaped, align='R')
        self.ln(1)

    def render_table(self, headers: list[str], rows: list[list[str]]):
        """Render an HTML table as a styled PDF table."""
        if not rows:
            return
        col_count = max(len(headers), max((len(r) for r in rows), default=0))
        if col_count == 0:
            return

        usable_w   = self.w - self.l_margin - self.r_margin
        col_w      = usable_w / col_count
        row_h      = 7
        self.set_arabic_font(9)

        # Header row
        self.set_fill_color(*self.INDIGO)
        self.set_text_color(*self.WHITE)
        self.set_draw_color(*self.BORDER)
        for h in (headers if headers else [""] * col_count):
            self.cell(col_w, row_h, shape_arabic(str(h)), border=1, fill=True, align='C')
        self.ln()

        # Data rows
        self.set_text_color(*self.DARK)
        for i, row in enumerate(rows):
            fill = i % 2 == 0
            self.set_fill_color(241, 245, 249) if fill else self.set_fill_color(*self.WHITE)
            for ci in range(col_count):
                cell_text = row[ci] if ci < len(row) else ""
                self.cell(col_w, row_h, shape_arabic(str(cell_text)), border=1, fill=fill, align='R')
            self.ln()
        self.ln(3)

def _extract_tables(html: str) -> list[dict]:
    """Return list of {headers, rows} dicts from HTML table tags."""
    tables_data = []
    table_blocks = re.findall(r'<table[^>]*>(.*?)</table>', html, re.DOTALL | re.IGNORECASE)
    for block in table_blocks:
        rows_html = re.findall(r'<tr[^>]*>(.*?)</tr>', block, re.DOTALL | re.IGNORECASE)
        headers, rows = [], []
        for ri, row_html in enumerate(rows_html):
            th_cells = re.findall(r'<th[^>]*>(.*?)</th>', row_html, re.DOTALL | re.IGNORECASE)
            td_cells = re.findall(r'<td[^>]*>(.*?)</td>', row_html, re.DOTALL | re.IGNORECASE)
            if th_cells:
                headers = [strip_tags(c).strip() for c in th_cells]
            elif td_cells:
                rows.append([strip_tags(c).strip() for c in td_cells])
        tables_data.append({"headers": headers, "rows": rows})
    return tables_data

@st.cache_data(show_spinner=False)
def create_pdf(content: str) -> bytes | None:
    """
    Generate a professional Arabic-ready PDF from extracted content.
    """
    import sys, io
    old_stdout = sys.stdout
    sys.stdout = io.StringIO()
    try:
        current_dir  = os.path.dirname(os.path.abspath(__file__))
        font_path    = os.path.join(current_dir, 'Arial.ttf')

        pdf = ArabicPDF(font_path=font_path)
        _ = pdf.add_page()

        cleaned = clean_html(content)
        pattern = re.compile(
            r'(<h[1-3][^>]*>.*?</h[1-3]>|<table[^>]*>.*?</table>)',
            re.DOTALL | re.IGNORECASE
        )
        segments = pattern.split(cleaned)

        for seg in segments:
            seg = seg.strip()
            if not seg:
                continue
            heading_match = re.match(r'<h[1-3][^>]*>(.*?)</h[1-3]>', seg, re.DOTALL | re.IGNORECASE)
            if heading_match:
                heading_text = strip_tags(heading_match.group(1)).strip()
                if heading_text:
                    _ = pdf.section_title(heading_text)
                continue
            if re.match(r'<table', seg, re.IGNORECASE):
                tables_data = _extract_tables(seg)
                for t in tables_data:
                    if t["headers"] or t["rows"]:
                        _ = pdf.render_table(t["headers"], t["rows"])
                continue
            plain = strip_tags(seg)
            for line in plain.split('\n'):
                line = line.strip()
                if line:
                    _ = pdf.body_line(line)
        out = pdf.output(dest='S')
        return out.encode('latin-1') if isinstance(out, str) else bytes(out)
    except Exception as e:
        sys.stdout = old_stdout
        st.warning(f"خطأ في إنشاء PDF: {e}")
        return None
    finally:
        sys.stdout = old_stdout

@st.cache_data(show_spinner=False)
def create_json_export(content: str, meta: dict) -> bytes:
    data = {
        "tool":       "Ink to text AI",
        "version":    "2.1",
        "timestamp":  datetime.now().isoformat(),
        "metadata":   meta,
        "raw_output": content,
        "plain_text": strip_tags(clean_html(content)),
        "word_count": count_words(content),
        "char_count": count_chars(content),
    }
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")

@st.cache_data(show_spinner=False)
def create_csv_from_tables(content: str) -> bytes | None:
    tables = re.findall(r'<table>(.*?)</table>', content, re.DOTALL | re.IGNORECASE)
    if not tables:
        return None
    csv_lines = []
    for t in tables:
        rows = re.findall(r'<tr>(.*?)</tr>', t, re.DOTALL | re.IGNORECASE)
        for row in rows:
            cells = re.findall(r'<(?:td|th)[^>]*>(.*?)</(?:td|th)>', row, re.DOTALL | re.IGNORECASE)
            csv_lines.append(",".join(f'"{strip_tags(c).strip()}"' for c in cells))
        csv_lines.append("")
    return "\n".join(csv_lines).encode("utf-8")

# ============================================================
def call_gemini(img_b64: str, prompt: str, temperature: float = 0.0) -> tuple[str, float]:
    payload = {
        "contents": [{
            "parts": [
                {"text": prompt},
                {"inlineData": {"mimeType": "image/jpeg", "data": img_b64}}
            ]
        }],
        "generationConfig": {
            "temperature":     temperature,
            "maxOutputTokens": 8192,
            "topP":            0.95,
        }
    }
    
    models_to_try = [
        "gemini-2.5-flash",
        "gemini-2.0-flash",
        "gemini-flash-latest"
    ]
    
    last_error = ""
    t0  = time.time()
    
    for model_name in models_to_try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={API_KEY}"
        try:
            res = requests.post(url, json=payload, timeout=30)
            if res.status_code == 200:
                text = res.json()['candidates'][0]['content']['parts'][0]['text']
                elapsed = float(f"{time.time() - t0:.2f}")
                return text, elapsed
            else:
                last_error = f"{res.status_code} ({model_name}): {res.text[:150]}"
                continue # حاول الموديل الذي بعده
        except Exception as e:
            last_error = f"Request Exception: {str(e)}"
            continue
            
    # إذا فشلت جميع الموديلات
    raise RuntimeError(f"All available models failed! Last seen error: {last_error}")

def build_prompt(mode: str, extra_instructions: str = "", lang_hint: str = "auto") -> str:
    lang_note = "" if lang_hint == "auto" else f"Output language preference: {lang_hint}. "
    base = (
        "You are a precision OCR engine. "
        "Extract ALL text and tables from this image exactly as they appear. "
        "Preserve original structure, formatting and layout. "
        "Render any table as valid HTML <table> with <tr> and <td> tags. "
        "CRITICAL: Do NOT wrap the output in markdown code blocks like ```html. Output raw plain text and tags ONLY! "
        "Do NOT add any commentary, explanation, or preamble — output raw content only. "
        f"{lang_note}"
    )
    modes = {
        "standard": base,
        "enhanced": (
            base +
            "Additionally: identify headings and format them with <h3> tags. "
            "Preserve paragraph breaks with <br>. "
            "Mark any handwritten text with [handwritten]. "
        ),
        "table_focus": (
            "You are a table-extraction specialist. "
            "Find and extract ALL tables in the image. "
            "Output only valid HTML tables — nothing else. "
            "Include headers where visible. "
            "If no tables are found, output: [No tables detected]. "
        ),
        "summary": (
            base +
            "After extraction, add a line break and then provide a concise 2-sentence summary "
            "prefixed with: SUMMARY: "
        ),
        "key_value": (
            "Extract all key-value pairs, form fields, or labeled data from this image. "
            "Format as an HTML table with two columns: Key | Value. "
            "If no structured pairs found, extract plain text normally. "
        ),
    }
    prompt = modes.get(mode, base)
    if extra_instructions.strip():
        prompt += f"\nAdditional instructions: {extra_instructions.strip()}"
    return prompt

# ============================================================
# 8. DEFAULTS & GLOBAL CONFIG ITEMS
# ============================================================
extraction_mode = "standard"
auto_enhance = True
contrast_val = 1.3
sharpness_val = 1.4
brightness_val = 1.05
denoise_val = False
api_temp = 0.0
output_lang = "auto"
extra_instr = ""
max_img_dim = 2048

# ============================================================
# 9. NAVIGATION & THEME TOGGLE
# ============================================================
# 9. THEME TOGGLE (Top Right)
# ============================================================
st.markdown('<div id="theme-btn-anchor"></div>', unsafe_allow_html=True)
_is_dark = st.session_state.get("dark_mode", False)
if _is_dark:
    theme_btn = st.button("☀️ Light Mode", key="theme_toggle")
else:
    theme_btn = st.button("🌙 Dark Mode", key="theme_toggle")

if theme_btn:
    st.session_state.dark_mode = not _is_dark
    st.rerun()

# ============================================================
# 10. MAIN NAVIGATION TABS
# ============================================================
tab_single, tab_batch, tab_history, tab_guide = st.tabs([
    "Single Document",
    "Batch Processing",
    "History",
    "How to Use"
])

# ============================================================
# TAB 1: SINGLE DOCUMENT
# ============================================================
with tab_single:
    st.markdown("""
<div class="simple-header">
    <h1>Ink to text AI</h1>
    <p>Convert handwritten notes to digital text. AI-powered recognition for messy handwriting.</p>
</div>
<div style="display:flex;justify-content:center;gap:1.5rem;font-size:0.95rem;font-weight:600;color:var(--text2);margin-bottom:2.5rem;margin-top:-1.5rem;">
    <span><i class="fa-solid fa-check" style="color:var(--success);"></i> Free to use</span>
    <span><i class="fa-solid fa-bolt" style="color:#f59e0b;"></i> ~10 seconds</span>
    <span><i class="fa-solid fa-copy" style="color:var(--accent);"></i> Export in seconds</span>
</div>
""", unsafe_allow_html=True)

    _u1, _u2, _u3 = st.columns([1, 8, 1])
    with _u2:
        uploaded_file = st.file_uploader(
            "Drop your image here",
            type=SUPPORTED_FORMATS,
            key="main_uploader",
            label_visibility="collapsed"
        )

        if uploaded_file:
            img_raw = ImageOps.exif_transpose(Image.open(uploaded_file))
            if auto_enhance:
                img_proc = enhance_image(
                    img_raw,
                    contrast=contrast_val,
                    sharpness=sharpness_val,
                    brightness=brightness_val,
                    denoise=denoise_val
                )
            else:
                img_proc = img_raw

            img_proc = resize_for_api(img_proc, max_px=max_img_dim)

            fsize = file_size_mb(uploaded_file)
            w, h  = img_raw.size
            st.session_state.current_image_meta = {
                "filename": uploaded_file.name,
                "size_mb":  fsize,
                "width":    w,
                "height":   h,
                "hash":     file_hash(uploaded_file),
                "mode":     extraction_mode,
            }

            st.image(img_proc, use_container_width=True, caption=uploaded_file.name)

            st.markdown("<br>", unsafe_allow_html=True)

            _b1, _b2, _b3 = st.columns([1,2,1])
            with _b2:
                extract_clicked = st.button("Extract Text Now", key="extract_btn", use_container_width=True)

            if extract_clicked:
                prompt = build_prompt(extraction_mode, extra_instr, output_lang)

                progress_placeholder = st.empty()
                progress_placeholder.markdown("""
                <div class="processing-container">
                    <div class="processing-spinner"></div>
                    <div class="processing-text">AI is deep-scanning your document...</div>
                </div>
                """, unsafe_allow_html=True)
                
                try:
                    img_b64 = image_to_b64(img_proc)
                    result, elapsed = call_gemini(img_b64, prompt, temperature=api_temp)
                    
                    st.session_state.extracted_result   = result
                    st.session_state.last_process_time  = elapsed
                    st.session_state.total_processed   += 1
                    st.session_state.total_chars       += count_chars(result)

                    st.session_state.processing_history.insert(0, {
                        "filename":  uploaded_file.name,
                        "timestamp": datetime.now().strftime("%H:%M:%S"),
                        "mode":      extraction_mode,
                        "words":     count_words(result),
                        "chars":     count_chars(result),
                        "lang":      detect_language(result),
                        "time":      elapsed,
                        "result":    result,
                    })

                    progress_placeholder.empty()
                    st.balloons()
                    st.toast(f"Extraction Successful in {elapsed}s", icon="✅")

                except Exception as e:
                    progress_placeholder.empty()
                    st.error(f"Extraction failed: {e}")
        else:
            pass # We removed the Awaiting Upload box since the dropzone itself is obvious

    # ============================================================
    # EXTRACTED CONTENT (Full Width Below)
    # ============================================================
    st.markdown("<hr style='margin:3rem 0; border:var(--border);'>", unsafe_allow_html=True)
    st.markdown(
        '<div class="card-title" style="justify-content:center;font-size:1.8rem;margin-bottom:2rem;"><i class="fa-solid fa-file-export" style="color:var(--accent);"></i> Extracted Content</div>',
        unsafe_allow_html=True
    )

    if st.session_state.extracted_result:
        result  = st.session_state.extracted_result
        display = clean_html(result)

        r1, r2, r3, r4 = st.columns(4)
        r1.metric("Words",    f"{count_words(result):,}")
        r2.metric("Chars",    f"{count_chars(result):,}")
        r3.metric("Language", detect_language(result))
        r4.metric("Type",     detect_content_type(result))

        st.markdown("<br>", unsafe_allow_html=True)

        view_tab, edit_tab = st.tabs(["Preview", "Edit & Notes"])

        with view_tab:
            preview_html = display.replace('\n', '<br>')
            st.markdown(f'<div class="output-box">{preview_html}</div>', unsafe_allow_html=True)

        with edit_tab:
            edited = st.text_area(
                "Edit extracted text:",
                value=strip_tags(display),
                height=300,
                label_visibility="collapsed"
            )
            st.session_state.notes = st.text_area(
                "Add notes:",
                value=st.session_state.notes,
                height=100,
                placeholder="Your personal notes about this document..."
            )
            if st.button("Save Edits", key="save_edit"):
                st.session_state.extracted_result = edited
                st.success("Edits saved!")

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown(
            '<div style="font-size:0.7rem;font-weight:600;color:#6b7280;margin-bottom:0.5rem;">'
            '<i class="fa-solid fa-box-archive" style="color:#4f46e5;margin-left:0.3rem;"></i>'
            ' Export Options'
            '</div>',
            unsafe_allow_html=True
        )

        export_ts = datetime.fromtimestamp(st.session_state.last_process_time).strftime('%Y%m%d_%H%M%S')
        base_name = f"DocuVision_{export_ts}"

        ec1, ec2, ec3 = st.columns(3)
        ec4, ec5, ec6 = st.columns(3)

        with ec1:
            st.download_button(
                "📄 Word Document",
                data=create_word_doc(result),
                file_name=f"{base_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_word_{export_ts}"
            )
        with ec2:
            pdf_bytes = create_pdf(result)
            if pdf_bytes:
                st.download_button(
                    "📕 PDF Document",
                    data=pdf_bytes,
                    file_name=f"{base_name}.pdf",
                    mime="application/pdf",
                    key=f"dl_pdf_{export_ts}"
                )
            else:
                st.button("📕 PDF (error)", disabled=True)
        with ec3:
            st.download_button(
                "📝 Plain Text",
                data=create_txt(result),
                file_name=f"{base_name}.txt",
                mime="text/plain",
                key=f"dl_txt_{export_ts}"
            )
        with ec4:
            st.download_button(
                "🌐 HTML Page",
                data=create_html_export(result),
                file_name=f"{base_name}.html",
                mime="text/html",
                key=f"dl_html_{export_ts}"
            )
        with ec5:
            st.download_button(
                "⚙️ JSON Export",
                data=create_json_export(result, st.session_state.current_image_meta),
                file_name=f"{base_name}.json",
                mime="application/json",
                key=f"dl_json_{export_ts}"
            )
        with ec6:
            csv_data = create_csv_from_tables(result)
            if csv_data:
                st.download_button(
                    "📊 CSV Tables",
                    data=csv_data,
                    file_name=f"{base_name}_tables.csv",
                    mime="text/csv",
                    key=f"dl_csv_{export_ts}"
                )
            else:
                st.button("📊 CSV (No tables)", disabled=True)

        with st.expander("Raw Output (for copy)", expanded=False):
            st.code(strip_tags(display), language=None)

    else:
        st.markdown("""
        <div style="text-align:center;padding:4rem 2rem;border:1px solid #e5e7eb;border-radius:1rem;">
            <i class="fa-solid fa-magnifying-glass"
               style="font-size:2.5rem;margin-bottom:0.5rem;display:block;color:#cbd5e1;"></i>
            <div style="font-weight:500;">Awaiting Extraction</div>
            <div style="font-size:0.8rem;color:#6b7280;">Upload an image and click Extract to see results here.</div>
        </div>
        <div class="steps-row" style="margin-top:1.5rem">
            <div class="step-item">
                <div class="step-num">1</div>
                <div><i class="fa-solid fa-upload icon-blue"></i> Upload</div>
                <div class="step-text">Drop your image</div>
            </div>
            <div class="step-item">
                <div class="step-num">2</div>
                <div><i class="fa-solid fa-sliders icon-blue"></i> Configure</div>
                <div class="step-text">Choose mode &amp; options</div>
            </div>
            <div class="step-item">
                <div class="step-num">3</div>
                <div><i class="fa-solid fa-robot icon-blue"></i> Extract</div>
                <div class="step-text">AI processes</div>
            </div>
            <div class="step-item">
                <div class="step-num">4</div>
                <div><i class="fa-solid fa-download icon-blue"></i> Export</div>
                <div class="step-text">Download in 6 formats</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

# ============================================================
# TAB 2: BATCH PROCESSING
# ============================================================
with tab_batch:
    st.markdown(
        '<div class="card-title"><i class="fa-solid fa-boxes-stacked"></i> Batch Document Processing</div>',
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='color:#6b7280;font-size:0.85rem;margin-bottom:1rem;'>"
        "Upload multiple images and extract text from all at once. Results are bundled into a ZIP file."
        "</p>",
        unsafe_allow_html=True
    )

    batch_files = st.file_uploader(
        "Upload multiple images",
        type=SUPPORTED_FORMATS,
        accept_multiple_files=True,
        key="batch_uploader",
        label_visibility="collapsed"
    )

    if batch_files:
        st.markdown(
            f"<div style='background:#f3f4f6;border-radius:0.75rem;padding:0.5rem 1rem;"
            f"margin-bottom:1rem;font-size:0.85rem;'>"
            f"<i class='fa-solid fa-folder-open' style='color:#4f46e5;'></i>"
            f" <strong>{len(batch_files)} files</strong> selected</div>",
            unsafe_allow_html=True
        )

        cols_preview = st.columns(min(len(batch_files), 4))
        for i, f in enumerate(batch_files[:4]):
            with cols_preview[i]:
                preview_img = ImageOps.exif_transpose(Image.open(f))
                st.image(preview_img, use_container_width=True, caption=f.name[:18])
        if len(batch_files) > 4:
            st.caption(f"... and {len(batch_files)-4} more files")

        st.markdown("<br>", unsafe_allow_html=True)
        bc1, bc2 = st.columns(2)
        with bc1:
            batch_mode = st.selectbox(
                "Extraction mode for all files",
                ["standard", "enhanced", "table_focus", "summary", "key_value"],
                format_func=lambda x: {
                    "standard":    "Standard",
                    "enhanced":    "Enhanced",
                    "table_focus": "Tables Only",
                    "summary":     "Extract + Summarize",
                    "key_value":   "Key-Value / Forms",
                }[x]
            )
        with bc2:
            batch_enhance = st.checkbox("Apply enhancement to all", value=True)

        if st.button("Start Batch Extraction", key="batch_btn"):
            batch_results  = []
            batch_progress = st.progress(0, text="Starting batch...")
            batch_status   = st.empty()
            total          = len(batch_files)

            for idx, bf in enumerate(batch_files):
                try:
                    batch_status.markdown(f"Processing **{bf.name}** ({idx+1}/{total})…")
                    img = ImageOps.exif_transpose(Image.open(bf))
                    if batch_enhance:
                        img = enhance_image(img, contrast=1.3, sharpness=1.4, brightness=1.05)
                    img = resize_for_api(img, max_px=2048)
                    b64 = image_to_b64(img)
                    prompt = build_prompt(batch_mode)
                    text, elapsed = call_gemini(b64, prompt)
                    batch_results.append({
                        "filename": bf.name,
                        "result":   text,
                        "words":    count_words(text),
                        "time":     elapsed,
                        "status":   "success"
                    })
                except Exception as e:
                    batch_results.append({
                        "filename": bf.name,
                        "result":   "",
                        "words":    0,
                        "time":     0,
                        "status":   f"error: {e}"
                    })
                batch_progress.progress((idx + 1) / total, text=f"Processed {idx+1}/{total}")

            batch_status.empty()
            batch_progress.empty()
            st.session_state.batch_results = batch_results

            success_count = sum(1 for r in batch_results if r["status"] == "success")
            st.success(f"Batch complete: {success_count}/{total} files processed successfully")

        if st.session_state.batch_results:
            st.markdown("---")
            st.markdown("### Batch Results")
            total_words = sum(r["words"] for r in st.session_state.batch_results)
            bm1, bm2, bm3 = st.columns(3)
            bm1.metric("Files Processed", len(st.session_state.batch_results))
            bm2.metric("Total Words", f"{total_words:,}")
            bm3.metric(
                "Success Rate",
                f"{sum(1 for r in st.session_state.batch_results if r['status']=='success')}"
                f"/{len(st.session_state.batch_results)}"
            )

            for r in st.session_state.batch_results:
                ok    = r["status"] == "success"
                icon  = "fa-circle-check icon-green" if ok else "fa-circle-xmark icon-red"
                color = "#10b981" if ok else "#ef4444"
                st.markdown(f"""
                <div class="history-row">
                    <i class="fa-solid {icon}"></i>
                    <span class="history-name">{r['filename']}</span>
                    <span style="color:{color};font-size:0.8rem;">{r['words']:,} words · {r['time']}s</span>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for r in st.session_state.batch_results:
                    if r["status"] == "success":
                        base_name = os.path.splitext(r["filename"])[0]
                        zf.writestr(f"{base_name}.txt", strip_tags(clean_html(r["result"])))
            zip_buf.seek(0)
            st.download_button(
                "Download All as ZIP (.txt files)",
                data=zip_buf.getvalue(),
                file_name=f"DocuVision_Batch_{datetime.now().strftime('%Y%m%d_%H%M')}.zip",
                mime="application/zip",
                key="dl_zip"
            )

# ============================================================
# TAB 3: HISTORY
# ============================================================
with tab_history:
    st.markdown(
        '<div class="card-title"><i class="fa-solid fa-clock-rotate-left"></i> Processing History</div>',
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='color:#6b7280;font-size:0.85rem;margin-bottom:1rem;'>"
        "Review all documents processed during this session."
        "</p>",
        unsafe_allow_html=True
    )

    if not st.session_state.processing_history:
        st.markdown(
            "<div style='text-align:center;padding:5rem 2rem;border:2px dashed var(--border2);border-radius:1.5rem;background:var(--bg2);'>"
            "<i class='fa-solid fa-clock-rotate-left' style='font-size:3.5rem;color:var(--border2);display:block;margin-bottom:1rem;'></i>"
            "<h3 style='margin:0 0 0.5rem 0;color:var(--text);font-size:1.3rem;font-weight:600;'>History is empty</h3>"
            "<p style='color:var(--text3);font-size:0.95rem;margin:0;'>Process your first document, and it will be safely logged right here.</p>"
            "</div>",
            unsafe_allow_html=True
        )
    else:
        st.markdown('<div class="history-search">', unsafe_allow_html=True)
        hist_search = st.text_input(
            "Search history by filename...",
            placeholder="🔍 Type a filename to filter...",
            label_visibility="collapsed"
        )
        st.markdown('</div><br>', unsafe_allow_html=True)
        
        filtered = (
            [h for h in st.session_state.processing_history if hist_search.lower() in h["filename"].lower()]
            if hist_search else st.session_state.processing_history
        )

        for i, h in enumerate(filtered):
            expander_title = f"📄 {h['filename']} &nbsp;•&nbsp; 🕒 {h['timestamp']} &nbsp;•&nbsp; 📝 {h['words']:,} words"
            with st.expander(expander_title, expanded=(i == 0)):
                hc1, hc2, hc3, hc4 = st.columns(4)
                hc1.metric("Words", f"{h['words']:,}")
                hc2.metric("Chars", f"{h['chars']:,}")
                hc3.metric("Language", h["lang"])
                hc4.metric("Speed", f"{h['time']}s")
                
                hist_html = clean_html(h["result"]).replace('\n', '<br>')
                st.markdown(f'<div class="output-box" style="margin-top:1rem;margin-bottom:1.5rem;">{hist_html}</div>', unsafe_allow_html=True)

                st.markdown('<div style="font-size:0.8rem;color:var(--text3);margin-bottom:0.8rem;font-weight:600;">EXPORT RECORD</div>', unsafe_allow_html=True)
                hd1, hd2, hd3, hd4 = st.columns([1, 1, 1, 1])
                with hd1:
                    st.download_button(
                        "📄 Word Doc",
                        data=create_word_doc(h["result"]),
                        file_name=f"{os.path.splitext(h['filename'])[0]}.docx",
                        key=f"hist_word_{i}"
                    )
                with hd2:
                    st.download_button(
                        "📝 Plain Text",
                        data=create_txt(h["result"]),
                        file_name=f"{os.path.splitext(h['filename'])[0]}.txt",
                        key=f"hist_txt_{i}"
                    )

        if len(filtered) > 1:
            all_zip = io.BytesIO()
            with zipfile.ZipFile(all_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for h in filtered:
                    base = os.path.splitext(h["filename"])[0]
                    zf.writestr(f"{base}.txt", strip_tags(clean_html(h["result"])))
            all_zip.seek(0)
            st.download_button(
                "Export All History as ZIP",
                data=all_zip.getvalue(),
                file_name=f"DocuVision_History_{datetime.now().strftime('%Y%m%d')}.zip",
                mime="application/zip",
                key="dl_history_zip"
            )

# ============================================================
# TAB 4: HOW TO USE
# ============================================================
with tab_guide:
    st.markdown(
        '<div class="card-title"><i class="fa-solid fa-book-open"></i> User Guide</div>',
        unsafe_allow_html=True
    )

    gc1, gc2 = st.columns(2)
    with gc1:
        with st.expander("Quick Start", expanded=True):
            st.markdown(
                "1. **Upload** your image (PNG, JPG, WEBP, etc.)\n"
                "2. **Choose mode** from the sidebar\n"
                "3. **Enable enhancement** if needed\n"
                "4. Click **Extract Text Now**\n"
                "5. **Download** in your preferred format"
            )
        with st.expander("Extraction Modes"):
            st.markdown(
                "| Mode | Best For |\n"
                "|------|----------|\n"
                "| **Standard** | General documents, letters |\n"
                "| **Enhanced** | Complex layouts with headings |\n"
                "| **Tables Only** | Spreadsheets, invoices |\n"
                "| **Extract + Summarize** | Long documents needing a brief overview |\n"
                "| **Key-Value / Forms** | Forms, IDs, receipts |"
            )
        with st.expander("Image Enhancement"):
            st.markdown(
                "- **Contrast**: Makes text stand out\n"
                "- **Sharpness**: Sharpens blurry text\n"
                "- **Brightness**: Lighten dark scans\n"
                "- **Denoise**: Removes speckles from old documents"
            )
    with gc2:
        with st.expander("Export Formats"):
            st.markdown(
                "- **Word (.docx)** – Editable\n"
                "- **PDF (عربي)** – Archiving with full Arabic RTL support\n"
                "- **Plain Text** – Copy-paste\n"
                "- **HTML** – Web embedding\n"
                "- **JSON** – System integration\n"
                "- **CSV** – Excel tables"
            )
        with st.expander("Language Support"):
            st.markdown(
                "Gemini 1.5 Flash supports Arabic (full RTL), English, French, "
                "German, Spanish, Chinese, Japanese, Korean, and 30+ more."
            )
        with st.expander("Pro Tips"):
            st.markdown(
                "- For **handwriting**, use Enhanced mode with high Sharpness\n"
                "- Use **Custom Instructions** to skip watermarks\n"
                "- **AI Temperature 0** gives most accurate results"
            )

# ============================================================
# 11. FOOTER
# ============================================================
st.markdown("""
<div class="app-footer">
    <i class="fa-solid fa-pen-nib"></i>
    <strong>Ink to text AI by Qays Hijjawi</strong>
    <i class="fa-solid fa-circle" style="font-size:0.4rem;vertical-align:middle;margin:0 0.3rem;"></i>
    Clean OCR tool
    <i class="fa-solid fa-circle" style="font-size:0.4rem;vertical-align:middle;margin:0 0.3rem;"></i>
    Powered by AI
    <br><span>&copy; 2026</span>
</div>
""", unsafe_allow_html=True)
